from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    from PIL import Image
except Exception:  # pragma: no cover - optional at runtime
    Image = None


BODY_FONT = "Carlito"
EA_FONT = "WenQuanYi Micro Hei"
MONO_FONT = "DejaVu Sans Mono"
ACCENT = RGBColor(46, 116, 181)
DARK_ACCENT = RGBColor(31, 77, 120)
INK = RGBColor(23, 32, 42)
MUTED = RGBColor(86, 97, 111)
TABLE_WIDTH_DXA = 9360


def set_ea_font(run, *, latin: str = BODY_FONT, east_asia: str = EA_FONT) -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def set_style_font(style, *, latin: str = BODY_FONT, east_asia: str = EA_FONT) -> None:
    style.font.name = latin
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style._element.rPr.rFonts.set(qn("w:ascii"), latin)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), latin)


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def new_decimal_numbering_id(document: Document) -> str:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = str((max(abstract_ids) + 1) if abstract_ids else 1)
    num_id = str((max(num_ids) + 1) if num_ids else 1)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_decimal_numbering(paragraph, num_id: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text: str, *, header: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    add_inline_runs(paragraph, text)
    for run in paragraph.runs:
        set_ea_font(run)
        run.font.size = Pt(9 if not header else 9.5)
        run.font.color.rgb = INK
        if header:
            run.bold = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_ea_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    set_style_font(title)
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, DARK_ACCENT, 8, 4),
    ]:
        style = styles[name]
        set_style_font(style)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        set_style_font(style)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1

    code = styles.add_style("H3D Code Block", 1)
    set_style_font(code, latin=MONO_FONT, east_asia=EA_FONT)
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor(34, 40, 49)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.08)
    code.paragraph_format.right_indent = Inches(0.08)

    caption = styles.add_style("H3D Caption", 1)
    set_style_font(caption)
    caption.font.size = Pt(9)
    caption.font.color.rgb = MUTED
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Harmonize3D Paper Deliverable")
    set_ea_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_number(footer)


def add_inline_runs(paragraph, text: str, *, code_mode: bool = False) -> None:
    if code_mode:
        run = paragraph.add_run(text)
        set_ea_font(run, latin=MONO_FONT, east_asia=EA_FONT)
        run.font.size = Pt(8.5)
        return

    pattern = re.compile(r"(`([^`]+)`)|(\*\*([^*]+)\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_ea_font(run)
            run.font.size = Pt(10.5)
            run.font.color.rgb = INK
        if match.group(2) is not None:
            run = paragraph.add_run(match.group(2))
            set_ea_font(run, latin=MONO_FONT, east_asia=EA_FONT)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(34, 40, 49)
        elif match.group(4) is not None:
            run = paragraph.add_run(match.group(4))
            set_ea_font(run)
            run.font.size = Pt(10.5)
            run.font.color.rgb = INK
            run.bold = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_ea_font(run)
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and set(stripped.replace("|", "").replace(" ", "")) <= {"-", ":"}


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def table_widths(col_count: int) -> list[int]:
    if col_count == 2:
        return [3300, 6060]
    if col_count == 3:
        return [2700, 3330, 3330]
    if col_count == 4:
        return [2700, 2500, 2060, 2100]
    if col_count == 5:
        return [2350, 1400, 2100, 1700, 1810]
    base = TABLE_WIDTH_DXA // col_count
    widths = [base] * col_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    normalized = [row + [""] * (col_count - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=col_count)
    table.style = "Table Grid"
    widths = table_widths(col_count)
    set_table_geometry(table, widths)
    repeat_header_row(table.rows[0])

    for r_idx, row in enumerate(normalized):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, header=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            if col_count >= 3 and c_idx in {2, 3}:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(document: Document, markdown_path: Path, line: str) -> None:
    match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if not match:
        return
    caption, rel = match.groups()
    image_path = (markdown_path.parent / rel).resolve()
    if not image_path.exists():
        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, f"[图片缺失：{image_path}]")
        return

    max_width = Inches(6.35)
    if Image is not None:
        with Image.open(image_path) as im:
            width_px, height_px = im.size
        if height_px > width_px * 1.25:
            max_width = Inches(4.8)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=max_width)

    if caption:
        cap = document.add_paragraph(style="H3D Caption")
        cap.paragraph_format.keep_together = True
        cap.add_run(caption)
        for run in cap.runs:
            set_ea_font(run)
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph(style="H3D Code Block")
    paragraph.paragraph_format.keep_together = True
    set_paragraph_shading(paragraph, "F6F8FA")
    text = "\n".join(lines)
    text = text.replace("/", "/\u200b").replace("_", "_\u200b")
    add_inline_runs(paragraph, text, code_mode=True)


def add_heading(document: Document, text: str, level: int) -> None:
    if level == 0:
        paragraph = document.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph = document.add_heading(level=min(level, 3))
    add_inline_runs(paragraph, text)
    for run in paragraph.runs:
        if level == 0:
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(11, 37, 69)
        elif level == 1:
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = ACCENT
        elif level == 2:
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = ACCENT
        else:
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = DARK_ACCENT


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    add_inline_runs(paragraph, text.strip())


def build_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    configure_document(document)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    current_numbering_id: str | None = None

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not line.strip():
            current_numbering_id = None
            i += 1
            continue

        if line.startswith("# "):
            current_numbering_id = None
            add_heading(document, line[2:].strip(), 0)
            i += 1
            continue
        if line.startswith("## "):
            current_numbering_id = None
            add_heading(document, line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith("### "):
            current_numbering_id = None
            add_heading(document, line[4:].strip(), 2)
            i += 1
            continue

        if line.strip().startswith("!["):
            current_numbering_id = None
            add_image(document, markdown_path, line)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            current_numbering_id = None
            table_rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(split_table_row(lines[i]))
                i += 1
            add_table(document, table_rows)
            continue

        bullet = re.match(r"^\s*-\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if bullet:
            current_numbering_id = None
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet.group(1).strip())
            i += 1
            continue
        if numbered:
            if current_numbering_id is None:
                current_numbering_id = new_decimal_numbering_id(document)
            paragraph = document.add_paragraph()
            apply_decimal_numbering(paragraph, current_numbering_id)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.1
            add_inline_runs(paragraph, numbered.group(1).strip())
            i += 1
            continue

        current_numbering_id = None
        add_body_paragraph(document, line.strip())
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: create_paper_docx.py input.md output.docx", file=sys.stderr)
        return 2
    build_docx(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
