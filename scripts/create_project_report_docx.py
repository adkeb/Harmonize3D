from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, *, name: str = "Calibri", east_asia: str = "Microsoft YaHei", size: int | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)


def set_paragraph_font(paragraph, *, size: int | None = None, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)
        run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def parse_inline(paragraph, text: str, *, code: bool = False) -> None:
    if code:
        run = paragraph.add_run(text)
        set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9)
        return

    pattern = re.compile(r"`([^`]+)`|\*\*([^*]+)\*\*")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run)
        if match.group(1) is not None:
            run = paragraph.add_run(match.group(1))
            set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9)
        elif match.group(2) is not None:
            run = paragraph.add_run(match.group(2))
            set_run_font(run)
            run.bold = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            parse_inline(paragraph, value.strip())
            for run in paragraph.runs:
                set_run_font(run, size=9 if row_index else 10)
                if row_index == 0:
                    run.bold = True
            if row_index == 0:
                shade_cell(cell, "D9EAF7")


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and set(stripped.replace("|", "").replace(" ", "")) <= {"-", ":"}


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [part.strip() for part in stripped.split("|")]


def add_image(document: Document, md_path: Path, line: str) -> None:
    match = re.match(r"!?\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if not match:
        return
    caption, rel = match.groups()
    image_path = (md_path.parent / rel).resolve()
    if not image_path.exists():
        paragraph = document.add_paragraph()
        parse_inline(paragraph, f"[图片缺失：{image_path}]")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.4))
    if caption:
        cap = document.add_paragraph()
        cap.alignment = 1
        r = cap.add_run(caption)
        set_run_font(r, size=9)
        r.italic = True


def build_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Calibri Light"
    styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Heading 3"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                paragraph = document.add_paragraph()
                paragraph.style = "No Spacing"
                parse_inline(paragraph, "\n".join(code_lines), code=True)
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            parse_inline(paragraph, line[2:].strip())
            set_paragraph_font(paragraph, size=22, bold=True)
            add_page_break(document)
            i += 1
            continue

        if line.startswith("## "):
            paragraph = document.add_heading(level=1)
            parse_inline(paragraph, line[3:].strip())
            for run in paragraph.runs:
                set_run_font(run, size=16)
                run.bold = True
                run.font.color.rgb = RGBColor(31, 56, 100)
            i += 1
            continue

        if line.startswith("### "):
            paragraph = document.add_heading(level=2)
            parse_inline(paragraph, line[4:].strip())
            for run in paragraph.runs:
                set_run_font(run, size=13)
                run.bold = True
            i += 1
            continue

        if line.strip().startswith("!["):
            add_image(document, markdown_path, line)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(split_table_row(lines[i]))
                i += 1
            add_table(document, table_rows)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            parse_inline(paragraph, line[2:].strip())
            i += 1
            continue

        paragraph = document.add_paragraph()
        parse_inline(paragraph, line.strip())
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.12
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: create_project_report_docx.py input.md output.docx", file=sys.stderr)
        return 2
    build_docx(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
